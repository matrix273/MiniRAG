import tempfile
from pathlib import Path
from docx import Document

from pageindex.parsers import ParseResult
from pageindex.parsers.docx_parser import parse_docx


def _make_docx(paragraphs: list[tuple[str, str]]) -> Path:
    """创建测试用 docx 文件。paragraphs: [(style, text), ...]"""
    doc = Document()
    for style, text in paragraphs:
        doc.add_paragraph(text, style=style)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return Path(tmp.name)


def test_parse_docx_headings():
    path = _make_docx([
        ("Heading 1", "Chapter 1"),
        ("Normal", "Some body text"),
        ("Heading 2", "Section 1.1"),
        ("Normal", "More text"),
    ])
    result = parse_docx(path)

    assert isinstance(result, ParseResult)
    assert len(result.sections) == 2
    assert result.sections[0]["title"] == "Chapter 1"
    assert result.sections[0]["level"] == 1
    assert result.sections[1]["title"] == "Section 1.1"
    assert result.sections[1]["level"] == 2
    assert "# Chapter 1" in result.content
    assert "## Section 1.1" in result.content
    assert "Some body text" in result.content


def test_parse_docx_tables():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    result = parse_docx(Path(tmp.name))

    assert "| Name |" in result.content
    assert "| A |" in result.content


def test_parse_docx_metadata():
    path = _make_docx([("Normal", "Hello")])
    result = parse_docx(path)
    assert "paragraph_count" in result.metadata
    assert result.metadata["paragraph_count"] == 1
